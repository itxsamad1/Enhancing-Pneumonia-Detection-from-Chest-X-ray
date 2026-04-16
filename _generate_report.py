"""
Generate professional Word document for the Pneumonia Detection research pipeline.
"""
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

BASE_DIR = Path(__file__).parent
GRAPHS_DIR = BASE_DIR / "graphs"
OUTPUT_PATH = BASE_DIR / "Pneumonia_Detection_Research_Report.docx"

# ─── Styling Helpers ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(table, border_color='2E4057'):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), border_color)
                tcBorders.append(el)
            tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
        run.font.size = Pt(13)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x04, 0x8A, 0x81)
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    return p

def add_graph(doc, filename, caption, width=Inches(5.5)):
    path = GRAPHS_DIR / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=width)
        add_caption(doc, caption)
    else:
        add_body(doc, f"[Graph not found: {filename}]", italic=True)

def make_header_row(table, headers, bg='1A1A2E', fg='FFFFFF'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)

def fill_table_row(table, row_idx, values, bold_first=False, bg=None):
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(10)
        if bold_first and i == 0:
            run.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bg:
            set_cell_bg(cell, bg)

# ─── Main Document Builder ───────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # ─── Title Page ──────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    run = title.add_run("Enhancing Pneumonia Detection from Chest X-ray Images\nUsing Image Preprocessing and Deep Learning")
    run.font.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(20)
    r = sub.add_run("ResNet-18 | Offline Dataset Augmentation | Anti-Overfitting Pipeline")
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x04, 0x8A, 0x81)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(30)
    r = meta.add_run("Abdul Samad\nUndergraduate Research Project\n2024-2025")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # ─── 1. Introduction ─────────────────────────────────────────
    add_heading(doc, "1. Introduction", 1)
    add_body(doc, (
        "Pneumonia is an acute respiratory infection that inflames the alveoli of one or both lungs. "
        "According to the World Health Organization, pneumonia accounts for approximately 14% of all deaths "
        "of children under five years old globally. Early and accurate diagnosis is critical for effective treatment. "
        "Chest X-rays are the most commonly used tool for pneumonia diagnosis; however, their interpretation "
        "requires specialist radiologists whose availability can be limited in resource-constrained environments."
    ))
    add_body(doc, (
        "This report documents the complete technical pipeline of a deep learning system designed to "
        "automatically classify chest X-ray images as Normal or Pneumonia. The pipeline spans dataset "
        "collection, augmentation, model training with explicit anti-overfitting regularizations, and "
        "web-based inference deployment. The work forms part of an ongoing research initiative and serves "
        "as the second iteration (RP2) in a planned series of architecture comparison experiments."
    ))

    # ─── 2. Dataset Collection ───────────────────────────────────
    add_heading(doc, "2. Dataset Collection", 1)
    add_body(doc, (
        "The primary training data was sourced from the Kaggle Chest X-Ray Images (Pneumonia) dataset, "
        "originally curated by Guangzhou Women and Children's Medical Center. All X-ray images were "
        "screened and graded by certified radiologists prior to their inclusion in the dataset."
    ))
    add_body(doc, (
        "The raw dataset was downloaded and reorganized into a unified directory using a custom preparation "
        "script (prepare_dataset.py), splitting all available images from the original train, validation, and "
        "test splits into a clean 80/20 train-validation partition."
    ))

    add_heading(doc, "2.1 Original Dataset Composition", 2)

    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl, ['Split', 'Normal', 'Pneumonia', 'Total'])
    fill_table_row(tbl, 1, ['Training', '1,266', '3,418', '4,684'], bold_first=True, bg='F0F4F8')
    fill_table_row(tbl, 2, ['Validation', '317', '855', '1,172'], bold_first=True)
    fill_table_row(tbl, 3, ['Total', '1,583', '4,273', '5,856'], bold_first=True, bg='D0D9E8')
    set_cell_border(tbl)
    add_caption(doc, "Table 1: Original dataset distribution before augmentation.")

    add_body(doc, (
        "A critical observation was the natural class imbalance: pneumonia cases outnumbered normal cases "
        "at a ratio of approximately 2.7:1. Without correction, this imbalance can bias model predictions "
        "toward the majority class (Pneumonia), resulting in poor specificity and high false positive rates."
    ))

    # ─── 3. Dataset Augmentation ─────────────────────────────────
    add_heading(doc, "3. Dataset Augmentation", 1)
    add_body(doc, (
        "To address both the limited scale of the original dataset and the class imbalance, an offline "
        "dataset augmentation pipeline was implemented in augment_dataset.py. Unlike real-time augmentation "
        "which applies transformations during batch loading, offline augmentation physically saves every "
        "augmented image to disk. This offers two key research advantages:"
    ))
    add_bullet(doc, "The expanded dataset becomes a fixed, reproducible artifact — every architecture trained on it receives exactly the same inputs.")
    add_bullet(doc, "Results across multiple model architectures (ResNet-18, ResNet-50, DenseNet, etc.) become perfectly comparable in future experiments.")

    add_heading(doc, "3.1 Class Balancing Strategy", 2)
    add_body(doc, (
        "An asymmetric augmentation factor was applied deliberately to restore class balance. "
        "The Normal (minority) class received a high augmentation multiplier, while the Pneumonia (majority) "
        "class received a lower multiplier. This approach naturally re-balanced the training distribution "
        "without discarding any original data."
    ))

    tbl2 = doc.add_table(rows=3, cols=4)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl2, ['Class', 'Augmentation Factor', 'Original Count', 'Expanded Count'])
    fill_table_row(tbl2, 1, ['Normal', '8x (Class Balancing)', '1,583', '14,247'], bold_first=True, bg='F0F4F8')
    fill_table_row(tbl2, 2, ['Pneumonia', '3x (Standard)', '4,273', '17,092'], bold_first=True)
    set_cell_border(tbl2)
    add_caption(doc, "Table 2: Asymmetric augmentation factors applied for class balancing.")

    add_heading(doc, "3.2 Augmentation Techniques", 2)
    add_body(doc, (
        "All eight techniques applied were specifically validated for medical chest X-ray imaging. "
        "Physiologically invalid transformations (e.g., vertical flips that would represent an upside-down "
        "patient, or extreme colour distortions that would erase grayscale lung tissue features) were "
        "explicitly excluded."
    ))

    tbl3 = doc.add_table(rows=9, cols=3)
    tbl3.style = 'Table Grid'
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl3, ['Technique', 'Parameters', 'Medical Justification'])
    aug_rows = [
        ('Horizontal Flip', 'p = 0.5', 'Lung anatomy is bilaterally symmetrical. A mirrored X-ray remains anatomically valid.'),
        ('Random Rotation', '+-15 degrees', 'Accounts for patient positional tilt during X-ray acquisition.'),
        ('Brightness Adjustment', '+/-20% factor', 'Simulates variable X-ray exposure settings across different machines.'),
        ('CLAHE Contrast', 'Clip limit: 1.5 - 3.0', 'Mimics the digital windowing used by radiologists to view lung tissue detail.'),
        ('Gaussian Noise', 'Std Dev: 5 - 15', 'Replicates "quantum mottle" or sensor granularity in lower-quality X-ray systems.'),
        ('Image Sharpening', '3x3 Laplacian kernel', 'Enhances edge definition to simulate high-resolution equipment scanning.'),
        ('Random Zoom', '5 - 20% centre crop', 'Accounts for variation in the patient-to-image-plate distance.'),
        ('Combined', '2-3 random techniques', 'Applies multiple techniques in sequence for maximum diversity.'),
    ]
    for i, r in enumerate(aug_rows):
        bg = 'F0F4F8' if i % 2 == 0 else None
        fill_table_row(tbl3, i + 1, r, bold_first=True, bg=bg)
    set_cell_border(tbl3)
    add_caption(doc, "Table 3: Augmentation techniques and their medical justifications.")

    add_heading(doc, "3.3 Final Augmented Dataset", 2)

    tbl4 = doc.add_table(rows=4, cols=4)
    tbl4.style = 'Table Grid'
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl4, ['Split', 'Normal', 'Pneumonia', 'Total'])
    fill_table_row(tbl4, 1, ['Training', '11,394', '13,672', '25,066'], bold_first=True, bg='F0F4F8')
    fill_table_row(tbl4, 2, ['Validation', '2,853', '3,420', '6,273'], bold_first=True)
    fill_table_row(tbl4, 3, ['Total', '14,247', '17,092', '31,339'], bold_first=True, bg='D0D9E8')
    set_cell_border(tbl4)
    add_caption(doc, "Table 4: Final augmented dataset. Class ratio improved from 1:2.7 to 1:1.2.")

    # ─── 4. Model Architecture ───────────────────────────────────
    add_heading(doc, "4. Model Architecture", 1)
    add_body(doc, (
        "The model used in this experiment is ResNet-18 (Residual Network), a convolutional neural network "
        "introduced by He et al. in 2015 and pre-trained on the ImageNet dataset comprising 1.2 million "
        "images across 1,000 classes. Transfer learning was applied: the convolutional backbone was "
        "initialized with ImageNet weights, and only the final classification head was re-trained for "
        "the target binary task (Normal vs. Pneumonia)."
    ))

    add_heading(doc, "4.1 Modified Classifier Head", 2)
    add_body(doc, (
        "The standard ResNet-18 final fully-connected layer (512 -> 1000 classes) was replaced with a "
        "custom Sequential block optimized for binary classification and overfitting prevention:"
    ))
    add_bullet(doc, "Dropout(p=0.3): Randomly zeroes out 30% of neurons during each training step, forcing the network to learn redundant representations rather than memorizing exact pixel patterns from training images.")
    add_bullet(doc, "Linear(512, 2): Maps the 512-dimensional feature vector to 2 output logits (Normal, Pneumonia) for binary sigmoid classification.")

    add_heading(doc, "4.2 Loss Function", 2)
    add_body(doc, (
        "CrossEntropyLoss with label_smoothing=0.1 was applied. Standard one-hot labels are binary "
        "(1.0 / 0.0). Label smoothing softens these to (0.9 / 0.1), preventing the model from becoming "
        "overconfident, especially on ambiguous or noisy X-ray images. This is a standard regularization "
        "technique proven effective in medical imaging tasks."
    ))

    # ─── 5. Training Configuration ───────────────────────────────
    add_heading(doc, "5. Training Configuration", 1)
    add_body(doc, (
        "Training was conducted on a system equipped with an NVIDIA GeForce RTX 5060 Ti GPU with "
        "16 GB of VRAM and approximately 4,600 CUDA cores running CUDA 12.8. Mixed precision training "
        "(torch.amp) was enabled, which reduced memory overhead and accelerated training throughput. "
        "The entire 17-epoch run completed in approximately 8.2 minutes."
    ))

    tbl5 = doc.add_table(rows=11, cols=2)
    tbl5.style = 'Table Grid'
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl5, ['Parameter', 'Value'])
    cfg_rows = [
        ('Model Architecture',      'ResNet-18 (ImageNet Pre-trained)'),
        ('Target Epochs',           '30'),
        ('Actual Epochs Completed', '17 (Early Stopping triggered)'),
        ('Best Performing Epoch',   '7'),
        ('Batch Size',              '64'),
        ('Initial Learning Rate',   '0.001'),
        ('Optimizer',               'Adam (weight_decay = 1e-4)'),
        ('LR Scheduler',            'CosineAnnealingWarmRestarts (T0=10, T_mult=2)'),
        ('Mixed Precision',         'torch.amp (CUDA AMP)'),
        ('GPU',                     'NVIDIA RTX 5060 Ti (16 GB VRAM, ~4,600 CUDA Cores)'),
    ]
    for i, r in enumerate(cfg_rows):
        bg = 'F0F4F8' if i % 2 == 0 else None
        fill_table_row(tbl5, i+1, r, bold_first=True, bg=bg)
    set_cell_border(tbl5)
    add_caption(doc, "Table 5: Training hyperparameter configuration.")

    add_heading(doc, "5.1 Why 30 Epochs were Set", 2)
    add_body(doc, (
        "The maximum number of epochs was configured at 30 to allow the model sufficient time to converge "
        "on a large augmented dataset of 25,066 training images. However, training was not expected to "
        "always complete all 30 epochs. Early Stopping was configured with a patience parameter of 10, "
        "meaning that if the validation loss did not decrease for 10 consecutive epochs, training would "
        "automatically terminate. This approach protects against overfitting while still providing "
        "an adequate upper bound for convergence."
    ))

    add_heading(doc, "5.2 Why Training Stopped at Epoch 17", 2)
    add_body(doc, (
        "The best validation loss of 0.2347 was recorded at Epoch 7. After this point, the validation "
        "loss did not improve for 10 consecutive epochs (Epochs 8 through 17). On Epoch 17, the early "
        "stopping condition was met and training was automatically halted. The model weights from Epoch 7 "
        "were preserved and saved as the final output."
    ))
    add_body(doc, (
        "This behavior is expected and desirable: the CosineAnnealing LR scheduler resets the learning "
        "rate at Epoch 11, causing the optimizer to briefly re-explore a broader parameter space. Despite "
        "this, the validation loss sequence from Epochs 7-17 failed to surpass the Epoch 7 benchmark, "
        "confirming that the model had effectively reached its optimal generalization point."
    ))

    add_heading(doc, "5.3 Anti-Overfitting Regularization Stack", 2)
    tbl6 = doc.add_table(rows=9, cols=3)
    tbl6.style = 'Table Grid'
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl6, ['Approach', 'Technique', 'Configuration'])
    reg_rows = [
        ('A', 'Offline Augmentation',     '31,339 images, 8 techniques'),
        ('B', 'Real-time Augmentation',    'RandomErasing, RandomAffine, ColorJitter'),
        ('C', 'Dropout',                   'p = 0.3 in classifier head'),
        ('D', 'Weight Decay (L2)',          'lambda = 1e-4 in Adam optimizer'),
        ('E', 'Early Stopping',             'Patience = 10 epochs on val_loss'),
        ('F', 'LR Scheduling',              'CosineAnnealingWarmRestarts (T0=10)'),
        ('G', 'Gradient Clipping',          'max_norm = 1.0'),
        ('H', 'Label Smoothing',            'epsilon = 0.1'),
    ]
    for i, r in enumerate(reg_rows):
        bg = 'F0F4F8' if i % 2 == 0 else None
        fill_table_row(tbl6, i+1, r, bold_first=True, bg=bg)
    set_cell_border(tbl6)
    add_caption(doc, "Table 6: Complete anti-overfitting regularization stack applied during training.")

    # ─── 6. Training Results ─────────────────────────────────────
    add_heading(doc, "6. Training Results", 1)
    add_heading(doc, "6.1 Per-Epoch Training Log", 2)

    with open(GRAPHS_DIR / "training_history.json") as f:
        hist = json.load(f)

    n = len(hist['train_loss'])
    tbl7 = doc.add_table(rows=n+1, cols=8)
    tbl7.style = 'Table Grid'
    tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl7, ['Epoch', 'Train Loss', 'Train Acc', 'Val Loss', 'Val Acc', 'Precision', 'Recall', 'F1'])
    for i in range(n):
        bg = 'E8F5E9' if i == 6 else ('F0F4F8' if i % 2 == 0 else None)  # Highlight epoch 7 (index 6)
        vals = [
            str(i+1),
            f"{hist['train_loss'][i]:.4f}",
            f"{hist['train_acc'][i]:.4f}",
            f"{hist['val_loss'][i]:.4f}",
            f"{hist['val_acc'][i]:.4f}",
            f"{hist['precision'][i]:.4f}",
            f"{hist['recall'][i]:.4f}",
            f"{hist['f1'][i]:.4f}",
        ]
        fill_table_row(tbl7, i+1, vals, bg=bg)
    set_cell_border(tbl7)
    add_caption(doc, "Table 7: Per-epoch training metrics. Epoch 7 (highlighted) recorded the best validation performance.")

    add_heading(doc, "6.2 Best Epoch Performance Summary", 2)
    tbl8 = doc.add_table(rows=7, cols=2)
    tbl8.style = 'Table Grid'
    tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl8, ['Metric', 'Value (Epoch 7)'])
    perf_rows = [
        ('Validation Accuracy', '97.91%'),
        ('Validation Loss',     '0.2347'),
        ('Precision',           '98.24%'),
        ('Recall (Sensitivity)', '97.92%'),
        ('F1 Score',            '98.08%'),
        ('Specificity',         '97.90%'),
    ]
    for i, r in enumerate(perf_rows):
        bg = 'F0F4F8' if i % 2 == 0 else None
        fill_table_row(tbl8, i+1, r, bold_first=True, bg=bg)
    set_cell_border(tbl8)
    add_caption(doc, "Table 8: Best epoch performance metrics recorded at Epoch 7.")

    add_heading(doc, "6.3 Overfitting Analysis", 2)
    add_body(doc, (
        "A key indicator of overfitting is the gap between training accuracy and validation accuracy. "
        "At Epoch 7 (the best epoch), the training accuracy was 97.88% and the validation accuracy was "
        "97.91%, yielding an accuracy gap of -0.03%. This is essentially zero, directly confirming that "
        "the combined regularization stack successfully prevented overfitting. The model learned to "
        "generalize to unseen X-rays rather than memorizing the training data."
    ))

    # ─── 7. Performance Graphs ───────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "7. Training Performance Graphs", 1)
    add_body(doc, (
        "The following graphs were automatically generated at the conclusion of training. "
        "Each graph captures the evolution of a specific metric across all 17 training epochs."
    ))

    add_heading(doc, "7.1 Accuracy", 2)
    add_graph(doc, "accuracy_plot.png", "Figure 1: Training and Validation Accuracy over 17 epochs.")

    add_heading(doc, "7.2 Loss", 2)
    add_graph(doc, "loss_plot.png", "Figure 2: Training and Validation Loss over 17 epochs.")

    add_heading(doc, "7.3 Precision", 2)
    add_graph(doc, "precision_plot.png", "Figure 3: Validation Precision over 17 epochs.")

    add_heading(doc, "7.4 Recall", 2)
    add_graph(doc, "recall_plot.png", "Figure 4: Validation Recall (Sensitivity) over 17 epochs.")

    add_heading(doc, "7.5 F1 Score", 2)
    add_graph(doc, "f1_score_plot.png", "Figure 5: F1 Score over 17 epochs.")

    add_heading(doc, "7.6 Specificity", 2)
    add_graph(doc, "specificity_plot.png", "Figure 6: Specificity over 17 epochs.")

    add_heading(doc, "7.7 Learning Rate Schedule", 2)
    add_graph(doc, "learning_rate_plot.png", "Figure 7: Learning rate decay using CosineAnnealingWarmRestarts.")

    add_heading(doc, "7.8 Confusion Matrix", 2)
    add_graph(doc, "confusion_matrix.png", "Figure 8: Confusion Matrix on the Validation Set (Best Epoch 7).", width=Inches(4.0))

    # ─── 8. Inference Application ────────────────────────────────
    add_heading(doc, "8. Inference & Deployment", 1)
    add_body(doc, (
        "A Streamlit-based web application (app.py) was developed to provide an interactive inference "
        "interface. The application allows a user to:"
    ))
    add_bullet(doc, "Upload any chest X-ray image in JPG or PNG format.")
    add_bullet(doc, "Optionally apply one or more preprocessing filters (CLAHE, Histogram Equalization, Denoising, Image Sharpening) before classification.")
    add_bullet(doc, "Receive a real-time binary classification result (Normal or Pneumonia) with a confidence percentage.")
    add_bullet(doc, "Inspect the Grad-CAM (Gradient-weighted Class Activation Mapping) heatmap, which highlights the exact spatial regions of the X-ray that most strongly influenced the model's prediction.")
    add_body(doc, (
        "Grad-CAM provides a critical layer of clinical interpretability. Rather than producing a black-box "
        "decision, the model's attention regions can be visually compared to the anatomical locations a "
        "qualified radiologist would examine when diagnosing pneumonia, providing a validation mechanism "
        "beyond pure classification accuracy."
    ))

    # ─── 9. Conclusion ───────────────────────────────────────────
    add_heading(doc, "9. Conclusion", 1)
    add_body(doc, (
        "This experiment demonstrates that a ResNet-18 architecture combined with an explicit "
        "anti-overfitting regularization stack and an offline class-balanced augmentation pipeline can "
        "achieve 97.91% validation accuracy on chest X-ray pneumonia detection. Critically, this "
        "performance was achieved with zero observed overfitting (train-validation accuracy gap of -0.03%) "
        "and in under 9 minutes of training time on a consumer-grade GPU."
    ))
    add_body(doc, (
        "The augmented dataset (31,339 images) established in this phase will serve as the fixed "
        "experimental baseline for a second planned research paper comparing multiple model architectures "
        "(ResNet-50, DenseNet-121, EfficientNet, VGG-16) on identical data, "
        "enabling scientifically rigorous performance comparisons."
    ))

    doc.save(str(OUTPUT_PATH))
    print(f"Document saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
